import re
import smtplib
from email import encoders
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.oxml import OxmlElement
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from datetime import date, timedelta
from docx.enum.style import WD_STYLE_TYPE
from email.mime.multipart import MIMEMultipart


def count_weekdays(start_date, end_date):
    """Conta os dias úteis entre duas datas."""
    day_count = 0
    current_date = start_date
    while current_date <= end_date:
        if current_date.weekday() < 5:  # 0 = segunda, 4 = sexta
            day_count += 1
        current_date += timedelta(days=1)
    return day_count

def get_sprint_number(today):
    """
    Calcula o número da sprint com base na data atual.
    Cada sprint tem 15 dias úteis (ignora fins de semana).
    A Sprint 0 começa em 07/04/2025.
    """
    sprint_start_date = date(2025, 4, 7)
    
    if today < sprint_start_date:
        return -1  # antes da primeira sprint
    
    # Contar quantos dias úteis se passaram
    total_weekdays = count_weekdays(sprint_start_date, today)
    
    # Cada sprint tem 15 dias úteis
    sprint_number = total_weekdays // 15
    return sprint_number

def adicionar_bloco_codigo(doc, codigo, linguagem=""):
    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run(codigo)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')

    # Estilo visual (opcional)
    paragrafo_format = paragrafo.paragraph_format
    paragrafo_format.space_before = Pt(6)
    paragrafo_format.space_after = Pt(6)

    # Borda estilo "caixa de código"
    p = paragrafo._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '2')
        border.set(qn('w:color'), 'auto')
        pBdr.append(border)
    pPr.append(pBdr)
    
def processar_imagem(doc, markdown_image):
    # Regex melhorado para capturar casos incompletos
    match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', markdown_image)
    if match:
        try:
            img_path = match.group(2).strip()
            print(f"Tentando inserir imagem: {img_path}")
            
            # Adiciona a imagem com tamanho proporcional
            doc.add_picture(img_path, width=Inches(4))  # 4 polegadas = ~10cm
            
            # Adiciona legenda se existir
            if match.group(1).strip():
                p = doc.add_paragraph(style='Caption')
                p.add_run(f"Figura: {match.group(1).strip()}").italic = True
                
        except Exception as e:
            print(f"Erro ao inserir imagem '{img_path}': {type(e).__name__} - {str(e)}")
    
def processar_linha_markdown(doc, line):
    if 'Code' not in doc.styles:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(10)
    if 'Caption' not in doc.styles:
        caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = 'Arial'
        caption_style.font.size = Pt(10)
        caption_style.font.italic = True
        caption_style.paragraph_format.space_after = Pt(6)
    
    line = line.strip()
    if not line:
        return
    line = line.strip()
    if not line:
        return

    p = doc.add_paragraph()
    
    # Remove marcador de lista se existir
    if line.startswith('- '):
        line = line[2:]

    # Processa elementos especiais em sequência
    segments = re.split(r'(```|\*\*|\*|`|!\[)', line)
    in_bold = False
    in_italic = False
    in_code = False
    in_code_block = False
    image_buffer = None

    for seg in segments:
        if seg == '**':
            in_bold = not in_bold
        elif seg == '*':
            in_italic = not in_italic
        elif seg == '`':
            if not in_code_block:
                in_code = not in_code
        elif seg == '```':
            in_code_block = not in_code_block
        elif seg == '![':
            image_buffer = seg  # Inicia o buffer de imagem
        elif image_buffer is not None:
            image_buffer += seg  # Acumula partes da imagem
            
            # Verifica se fechou com ')' e processa
            if ')' in seg:
                # Processa mesmo que incompleto
                processar_imagem(doc, image_buffer)
                image_buffer = None
                p = doc.add_paragraph()  # Nova linha após imagem
        else:
            if seg:
                run = p.add_run(seg)
                
                # Aplica formatação hierarquicamente
                if in_code_block:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif in_code:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.style = 'Code'
                else:
                    run.bold = in_bold
                    run.italic = in_italic
    
def substituir_variaveis_em_documento(doc, variaveis):
    # Substitui em parágrafos fora de tabela
    for paragrafo in doc.paragraphs:
        for chave, valor in variaveis.items():
            if f"{{{{{chave}}}}}" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(f"{{{{{chave}}}}}", valor)

    # Substitui dentro de tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for chave, valor in variaveis.items():
                        if f"{{{{{chave}}}}}" in paragrafo.text:
                            paragrafo.text = paragrafo.text.replace(f"{{{{{chave}}}}}", valor)

def inserir_resposta_llm(doc, texto):
    # Primeiro processa blocos de código
    padrao_bloco_codigo = re.compile(r'```(\w+)?\n(.*?)\n```', re.DOTALL)
    pos = 0
    
    for match in padrao_bloco_codigo.finditer(texto):
        # Processa texto antes do bloco
        texto_anterior = texto[pos:match.start()].strip()
        if texto_anterior:
            for linha in texto_anterior.split('\n'):
                processar_linha_markdown(doc, linha)
        
        # Adiciona bloco de código
        adicionar_bloco_codigo(doc, match.group(2).strip(), match.group(1))
        pos = match.end()
    
    # Processa texto após o último bloco
    texto_restante = texto[pos:].strip()
    if texto_restante:
        for linha in texto_restante.split('\n'):
            processar_linha_markdown(doc, linha)
                
def enviar_email(pdf, dia):
    password = "ojbs qfgx euob kujr"
    sender = "alexandre.tavares@kuaracapital.com"
    destination = ["manuel@kuaracapital.com", "kaue.figueiredo@kuaracapital.com"]
    
    subject = f"[INTERNO] ATA do dia {dia}"
    body = f"""
    Prezados, bom dia.
    
    Segue em anexo a ata das atividades desenvolvidas no dia {dia}.
    
    Fico à disposição para quaisquer esclarecimentos.
    
    Atenciosamente, Alexandre Tavares.
    """
    
    msg = MIMEMultipart()
    msg['From'] = sender
    msg['To'] = ", ".join(destination) if isinstance(destination, list) else destination
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))
    with open(pdf, 'rb') as anexo:
        parte = MIMEBase('application', 'octet-stream')
        parte.set_payload(anexo.read())
        encoders.encode_base64(parte)
        parte.add_header('Content-Disposition', f'attachment; filename="{pdf}"')
        msg.attach(parte)
        
    try:
        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(sender, password)
            server.sendmail(sender, destination, msg.as_string())
        print("Email enviado com sucesso!")
    except Exception as e:
        print(f"Erro ao enviar email: {e}")